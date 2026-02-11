"""Document processor: extract text from PDF, DOCX, Excel, TXT."""

import mimetypes
from pathlib import Path
from typing import TYPE_CHECKING, Optional

from pypdf import PdfReader
from docx import Document as DocxDocument
import openpyxl

from .base import BaseDocumentProcessor
from .types import ExtractResult

if TYPE_CHECKING:
    from .ocr_processor import OcrProcessor


class DocumentProcessor(BaseDocumentProcessor):
    """
    Single entry for document text extraction. Supports PDF, DOCX, Excel, TXT.
    PDF behavior is pluggable: pass pdf_processor (e.g. OcrProcessor) to use
    PyMuPDF + pytesseract OCR fallback for PDFs; otherwise uses PyPDF (native text only).
    """

    def __init__(
        self,
        upload_dir: str | None = None,
        pdf_processor: Optional["OcrProcessor"] = None,
    ):
        if upload_dir is None:
            from ..config import get_settings
            upload_dir = get_settings().UPLOAD_DIR
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self._pdf_processor = pdf_processor

    def extract(self, file_path: str | Path) -> ExtractResult:
        """Extract text from a file. Returns ExtractResult with text and metadata."""
        path = Path(file_path)
        suffix = path.suffix.lower()
        try:
            if suffix == ".pdf":
                if self._pdf_processor is not None:
                    result = self._pdf_processor.extract(path)  # OcrProcessor handles PDF + OCR
                    # If OCR path failed (e.g. Tesseract not installed), try PyPDF for native text
                    if not (result.text or "").strip():
                        fallback = self._extract_pdf(path)
                        if (fallback.text or "").strip():
                            return fallback
                    return result
                return self._extract_pdf(path)
            if suffix in (".doc", ".docx"):
                return self._extract_docx(path)
            if suffix in (".xlsx", ".xls"):
                return self._extract_excel(path)
            if suffix == ".txt":
                return self._extract_txt(path)
            return ExtractResult.error_result(
                f"Unsupported type: {suffix}",
                metadata={"mime": mimetypes.guess_type(str(path))[0]},
            )
        except Exception as e:
            return ExtractResult.error_result(str(e))

    def _extract_pdf(self, path: Path) -> ExtractResult:
        reader = PdfReader(path)
        parts = [p.extract_text() or "" for p in reader.pages]
        meta = {}
        if reader.metadata:
            meta = {"title": reader.metadata.get("/Title"), "author": reader.metadata.get("/Author"), "pages": len(reader.pages)}
        return ExtractResult("\n\n".join(parts), metadata=meta)

    def _extract_docx(self, path: Path) -> ExtractResult:
        doc = DocxDocument(path)
        paragraphs = [p.text for p in doc.paragraphs]
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                tables_text.append(" | ".join(cell.text for cell in row.cells))
        text = "\n".join(paragraphs)
        if tables_text:
            text += "\n\n" + "\n".join(tables_text)
        return ExtractResult(text, metadata={"paragraphs": len(paragraphs), "tables": len(doc.tables)})

    def _extract_excel(self, path: Path) -> ExtractResult:
        wb = openpyxl.load_workbook(path, data_only=True)
        parts = []
        for name in wb.sheetnames:
            sheet = wb[name]
            parts.append(f"\n--- Sheet: {name} ---\n")
            for row in sheet.iter_rows(values_only=True):
                parts.append(" | ".join(str(c) if c is not None else "" for c in row))
        return ExtractResult("\n".join(parts), metadata={"sheets": wb.sheetnames})

    def _extract_txt(self, path: Path) -> ExtractResult:
        for enc in ("utf-8", "latin-1"):
            try:
                text = path.read_text(encoding=enc)
                return ExtractResult(text, metadata={"encoding": enc})
            except UnicodeDecodeError:
                continue
        return ExtractResult.error_result("Could not decode text file")

    def save_upload(self, content: bytes, filename: str) -> Path:
        """Save uploaded bytes to upload_dir with optional deduplication."""
        target = self.upload_dir / filename
        n = 1
        while target.exists():
            stem, suf = target.stem, target.suffix
            target = self.upload_dir / f"{stem}_{n}{suf}"
            n += 1
        target.write_bytes(content)
        return target
